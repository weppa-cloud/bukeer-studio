from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import ListFlowable, ListItem, PageBreak, Paragraph, SimpleDocTemplate, Spacer


OUT_DIR = Path("ops/google-ads/colombiatours/2026-05-launch/google_ads_api_application")
DOCX_PATH = OUT_DIR / "ColombiaTours_Google_Ads_API_Standard_Access_Design.docx"
PDF_PATH = OUT_DIR / "ColombiaTours_Google_Ads_API_Standard_Access_Design.pdf"


def font(size=28, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except Exception:
            pass
    return ImageFont.load_default()


def draw_button(draw, box, text, fill, outline="#1f2937", text_fill="#111827"):
    draw.rounded_rectangle(box, radius=12, fill=fill, outline=outline, width=2)
    x1, y1, x2, y2 = box
    bbox = draw.textbbox((0, 0), text, font=font(20, True))
    draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y1 + (y2 - y1 - (bbox[3] - bbox[1])) / 2 - 2), text, font=font(20, True), fill=text_fill)


def make_dashboard(path):
    img = Image.new("RGB", (1400, 850), "#f7f8fb")
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, 1400, 90), fill="#0f172a")
    d.text((40, 28), "ColombiaTours Ads Operations", font=font(34, True), fill="white")
    d.text((1050, 34), "Internal tool", font=font(24), fill="#cbd5e1")
    d.text((40, 125), "Campaign Performance Dashboard", font=font(38, True), fill="#111827")
    d.text((40, 175), "Google Ads API reporting for internal marketing review", font=font(24), fill="#4b5563")

    cards = [
        ("Impressions", "84,250", "#dbeafe"),
        ("Clicks", "5,430", "#dcfce7"),
        ("Conversions", "318", "#fef3c7"),
        ("Cost", "$1,038", "#fee2e2"),
    ]
    x = 40
    for title, value, color in cards:
        d.rounded_rectangle((x, 230, x + 300, 390), radius=16, fill=color, outline="#cbd5e1", width=2)
        d.text((x + 28, 255), title, font=font(24, True), fill="#111827")
        d.text((x + 28, 310), value, font=font(42, True), fill="#111827")
        x += 335

    d.rounded_rectangle((40, 440, 1360, 790), radius=16, fill="white", outline="#cbd5e1", width=2)
    d.text((70, 470), "Search campaigns", font=font(28, True), fill="#111827")
    headers = ["Campaign", "Status", "Clicks", "Cost", "Conversions"]
    widths = [520, 180, 180, 180, 220]
    y = 535
    x = 70
    for h, w in zip(headers, widths):
        d.rectangle((x, y, x + w, y + 50), fill="#e5e7eb", outline="#cbd5e1")
        d.text((x + 15, y + 13), h, font=font(20, True), fill="#111827")
        x += w
    rows = [
        ["MX_Multidestino_y_Caribe_2026_05", "Paused", "2,910", "$560", "184"],
        ["ES_Cartagena_Medellin_2026_05", "Paused", "2,520", "$478", "134"],
    ]
    y += 50
    for row in rows:
        x = 70
        for cell, w in zip(row, widths):
            d.rectangle((x, y, x + w, y + 58), fill="white", outline="#e5e7eb")
            d.text((x + 15, y + 17), cell, font=font(19), fill="#111827")
            x += w
        y += 58
    img.save(path)


def make_validator(path):
    img = Image.new("RGB", (1400, 850), "#f8fafc")
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, 1400, 90), fill="#0f766e")
    d.text((40, 28), "Campaign Upload Validator", font=font(34, True), fill="white")
    d.text((40, 125), "Dry-run validation before Google Ads API mutation", font=font(34, True), fill="#111827")
    d.rounded_rectangle((40, 195, 1360, 750), radius=16, fill="white", outline="#cbd5e1", width=2)
    checks = [
        ("Campaign names are unique", "PASS"),
        ("Required EU political ads field completed", "PASS"),
        ("Ad groups linked to existing campaigns", "PASS"),
        ("Responsive search ads have 15 headlines and 4 descriptions", "PASS"),
        ("Negative keywords loaded", "PASS"),
        ("Campaigns will be created paused", "PASS"),
    ]
    y = 245
    for label, status in checks:
        d.ellipse((85, y + 5, 125, y + 45), fill="#22c55e")
        d.text((96, y + 9), "✓", font=font(26, True), fill="white")
        d.text((150, y + 10), label, font=font(25), fill="#111827")
        draw_button(d, (1100, y, 1260, y + 50), status, "#dcfce7", "#22c55e", "#166534")
        y += 75
    draw_button(d, (70, 660, 350, 720), "Run dry-run", "#e0f2fe", "#0284c7", "#075985")
    draw_button(d, (380, 660, 680, 720), "Export review", "#fef3c7", "#d97706", "#92400e")
    img.save(path)


def make_review(path):
    img = Image.new("RGB", (1400, 850), "#f7f8fb")
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, 1400, 90), fill="#1d4ed8")
    d.text((40, 28), "Change Review Queue", font=font(34, True), fill="white")
    d.text((40, 125), "Human review before campaign launch", font=font(36, True), fill="#111827")
    d.rounded_rectangle((40, 190, 1360, 705), radius=16, fill="white", outline="#cbd5e1", width=2)
    items = [
        ("Campaigns", "2 new Search campaigns", "Paused"),
        ("Ad groups", "7 new ad groups", "Paused"),
        ("Keywords", "45 keywords", "Ready"),
        ("Responsive search ads", "7 ads", "Ready"),
        ("Negative keywords", "74 negatives", "Ready"),
        ("Assets", "Sitelinks, callouts, snippets", "Ready"),
    ]
    y = 240
    for title, detail, status in items:
        d.rounded_rectangle((80, y, 1320, y + 62), radius=10, fill="#f9fafb", outline="#e5e7eb")
        d.text((110, y + 16), title, font=font(24, True), fill="#111827")
        d.text((440, y + 17), detail, font=font(23), fill="#374151")
        d.text((1100, y + 17), status, font=font(23, True), fill="#0f766e")
        y += 72
    draw_button(d, (80, 750, 360, 815), "Create paused", "#dbeafe", "#2563eb", "#1e3a8a")
    draw_button(d, (390, 750, 690, 815), "Manual review", "#fef9c3", "#ca8a04", "#854d0e")
    draw_button(d, (720, 750, 1010, 815), "Enable in Ads UI", "#dcfce7", "#16a34a", "#166534")
    img.save(path)


def set_doc_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    for style_name, size in [("Title", 20), ("Heading 1", 15), ("Heading 2", 13)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)


def add_label_para(doc, label, text):
    p = doc.add_paragraph()
    run = p.add_run(label + ": ")
    run.bold = True
    p.add_run(text)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    dash = OUT_DIR / "mockup_1_performance_dashboard.png"
    validator = OUT_DIR / "mockup_2_campaign_validator.png"
    review = OUT_DIR / "mockup_3_change_review.png"
    make_dashboard(dash)
    make_validator(validator)
    make_review(review)

    doc = Document()
    set_doc_styles(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Google Ads API Standard Access Design Documentation")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(17, 24, 39)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("ColombiaTours Travel internal campaign operations tool")

    doc.add_paragraph()
    add_label_para(doc, "Company Name", "ColombiaTours Travel")
    add_label_para(doc, "Primary Website", "https://colombiatours.travel/")
    add_label_para(doc, "Tool Availability", "Internal users only. The tool is not publicly accessible and external users or clients will not have direct access.")

    doc.add_heading("Business Model", level=1)
    doc.add_paragraph(
        "ColombiaTours Travel is a travel agency that sells customized travel packages and tours in Colombia to international travelers. "
        "The company generates leads through its website and sales team, then helps travelers plan and purchase destination-specific travel packages. "
        "We advertise only our own travel services and websites, and we do not manage Google Ads accounts for third-party clients."
    )

    doc.add_heading("Tool Access and Use", level=1)
    doc.add_paragraph(
        "The Google Ads API tool will be used only by internal employees and managers of ColombiaTours Travel. "
        "The tool will help the internal marketing team create, validate, pause, update, and review Google Ads Search campaigns. "
        "Campaigns will be created from approved internal campaign templates and CSV bundles. All new campaigns, ad groups, keywords, and ads will be created in paused status first, reviewed internally, and then manually enabled by an authorized account user."
    )

    doc.add_heading("Tool Design", level=1)
    doc.add_paragraph(
        "The tool will use the Google Ads API to manage Search campaign structures for ColombiaTours Travel. Internal users will prepare approved campaign data such as campaign names, budgets, ad groups, keywords, negative keywords, responsive search ads, sitelinks, callouts, and structured snippets."
    )
    doc.add_paragraph(
        "Before applying changes, the tool will validate the campaign data, check required fields, verify naming conventions, and prevent duplicate campaign creation. The tool will support dry-run or validation mode before creating or updating account entities."
    )
    doc.add_paragraph(
        "The tool will also retrieve campaign performance metrics for internal reporting, including impressions, clicks, cost, conversions, CTR, CPC, and conversion rate. Reports will be used by the internal marketing team to review performance and optimize ColombiaTours campaigns."
    )

    doc.add_heading("Security and Access Control", level=1)
    doc.add_paragraph(
        "Only authorized internal users will be able to run the tool. API credentials and OAuth tokens will be stored securely and will not be exposed to end users. The developer token will not be shared with external users. The tool will only access Google Ads accounts owned or managed by ColombiaTours Travel."
    )

    doc.add_heading("API Services Called", level=1)
    services = [
        "GoogleAdsService: retrieve account, campaign, ad group, keyword, ad, and performance reporting data.",
        "CustomerService: access basic customer account information.",
        "CampaignBudgetService: create and update campaign budgets.",
        "CampaignService: create, pause, and update Search campaigns.",
        "AdGroupService: create, pause, and update ad groups.",
        "AdGroupCriterionService: create and update keywords and negative keywords.",
        "AdGroupAdService: create and update responsive search ads.",
        "AssetService: create ad assets such as sitelinks, callouts, and structured snippets.",
        "CampaignAssetService: attach campaign-level assets.",
        "CustomerNegativeCriterionService: manage account-level negative keywords when needed.",
    ]
    for item in services:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Supported Campaign Types", level=1)
    doc.add_paragraph("Search campaigns.")

    doc.add_heading("Primary Capabilities", level=1)
    for item in [
        "Campaign creation",
        "Campaign management",
        "Keyword management",
        "Negative keyword management",
        "Responsive search ad creation",
        "Ad asset management",
        "Reporting",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Change Review Process", level=1)
    doc.add_paragraph(
        "All bulk changes are reviewed before being applied. New campaigns are created in paused status by default. An authorized ColombiaTours user reviews the changes in Google Ads before enabling campaigns for delivery. The tool is intended to reduce manual upload errors and standardize campaign creation, not to automatically launch campaigns without human review."
    )

    doc.add_heading("Tool Mockups", level=1)
    for heading, image_path in [
        ("Mockup 1: Internal performance dashboard", dash),
        ("Mockup 2: Campaign upload dry-run validator", validator),
        ("Mockup 3: Human review queue before launch", review),
    ]:
        doc.add_heading(heading, level=2)
        doc.add_picture(str(image_path), width=Inches(6.5))

    doc.save(DOCX_PATH)
    build_pdf(dash, validator, review)
    print(DOCX_PATH)


def build_pdf(dash, validator, review):
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="DocTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=18, leading=22, textColor=colors.HexColor("#111827")))
    styles.add(ParagraphStyle(name="H1Custom", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=14, leading=18, spaceBefore=12, textColor=colors.HexColor("#111827")))
    styles.add(ParagraphStyle(name="BodyCustom", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=14, spaceAfter=8))
    styles.add(ParagraphStyle(name="Small", parent=styles["BodyText"], fontName="Helvetica", fontSize=9, leading=12, textColor=colors.HexColor("#4b5563")))

    story = [
        Paragraph("Google Ads API Standard Access Design Documentation", styles["DocTitle"]),
        Paragraph("ColombiaTours Travel internal campaign operations tool", styles["Small"]),
        Spacer(1, 0.18 * inch),
        Paragraph("<b>Company Name:</b> ColombiaTours Travel", styles["BodyCustom"]),
        Paragraph("<b>Primary Website:</b> https://colombiatours.travel/", styles["BodyCustom"]),
        Paragraph("<b>Tool Availability:</b> Internal users only. The tool is not publicly accessible and external users or clients will not have direct access.", styles["BodyCustom"]),
        Paragraph("Business Model", styles["H1Custom"]),
        Paragraph("ColombiaTours Travel is a travel agency that sells customized travel packages and tours in Colombia to international travelers. The company generates leads through its website and sales team, then helps travelers plan and purchase destination-specific travel packages. We advertise only our own travel services and websites, and we do not manage Google Ads accounts for third-party clients.", styles["BodyCustom"]),
        Paragraph("Tool Access and Use", styles["H1Custom"]),
        Paragraph("The Google Ads API tool will be used only by internal employees and managers of ColombiaTours Travel. The tool will help the internal marketing team create, validate, pause, update, and review Google Ads Search campaigns. Campaigns will be created from approved internal campaign templates and CSV bundles. All new campaigns, ad groups, keywords, and ads will be created in paused status first, reviewed internally, and then manually enabled by an authorized account user.", styles["BodyCustom"]),
        Paragraph("Tool Design", styles["H1Custom"]),
        Paragraph("The tool will use the Google Ads API to manage Search campaign structures for ColombiaTours Travel. Internal users will prepare approved campaign data such as campaign names, budgets, ad groups, keywords, negative keywords, responsive search ads, sitelinks, callouts, and structured snippets.", styles["BodyCustom"]),
        Paragraph("Before applying changes, the tool will validate the campaign data, check required fields, verify naming conventions, and prevent duplicate campaign creation. The tool will support dry-run or validation mode before creating or updating account entities.", styles["BodyCustom"]),
        Paragraph("The tool will also retrieve campaign performance metrics for internal reporting, including impressions, clicks, cost, conversions, CTR, CPC, and conversion rate. Reports will be used by the internal marketing team to review performance and optimize ColombiaTours campaigns.", styles["BodyCustom"]),
        Paragraph("Security and Access Control", styles["H1Custom"]),
        Paragraph("Only authorized internal users will be able to run the tool. API credentials and OAuth tokens will be stored securely and will not be exposed to end users. The developer token will not be shared with external users. The tool will only access Google Ads accounts owned or managed by ColombiaTours Travel.", styles["BodyCustom"]),
        Paragraph("API Services Called", styles["H1Custom"]),
    ]
    services = [
        "GoogleAdsService: retrieve account, campaign, ad group, keyword, ad, and performance reporting data.",
        "CustomerService: access basic customer account information.",
        "CampaignBudgetService: create and update campaign budgets.",
        "CampaignService: create, pause, and update Search campaigns.",
        "AdGroupService: create, pause, and update ad groups.",
        "AdGroupCriterionService: create and update keywords and negative keywords.",
        "AdGroupAdService: create and update responsive search ads.",
        "AssetService: create ad assets such as sitelinks, callouts, and structured snippets.",
        "CampaignAssetService: attach campaign-level assets.",
        "CustomerNegativeCriterionService: manage account-level negative keywords when needed.",
    ]
    story.append(ListFlowable([ListItem(Paragraph(item, styles["BodyCustom"])) for item in services], bulletType="bullet", leftIndent=18))
    story.extend([
        Paragraph("Supported Campaign Types", styles["H1Custom"]),
        Paragraph("Search campaigns.", styles["BodyCustom"]),
        Paragraph("Primary Capabilities", styles["H1Custom"]),
        ListFlowable([ListItem(Paragraph(item, styles["BodyCustom"])) for item in [
            "Campaign creation", "Campaign management", "Keyword management", "Negative keyword management",
            "Responsive search ad creation", "Ad asset management", "Reporting"
        ]], bulletType="bullet", leftIndent=18),
        Paragraph("Change Review Process", styles["H1Custom"]),
        Paragraph("All bulk changes are reviewed before being applied. New campaigns are created in paused status by default. An authorized ColombiaTours user reviews the changes in Google Ads before enabling campaigns for delivery. The tool is intended to reduce manual upload errors and standardize campaign creation, not to automatically launch campaigns without human review.", styles["BodyCustom"]),
        PageBreak(),
        Paragraph("Tool Mockups", styles["H1Custom"]),
    ])
    for heading, image_path in [
        ("Mockup 1: Internal performance dashboard", dash),
        ("Mockup 2: Campaign upload dry-run validator", validator),
        ("Mockup 3: Human review queue before launch", review),
    ]:
        story.append(Paragraph(heading, styles["Heading2"]))
        story.append(PdfImage(str(image_path), width=6.7 * inch, height=4.07 * inch))
        story.append(Spacer(1, 0.16 * inch))

    pdf = SimpleDocTemplate(str(PDF_PATH), pagesize=letter, rightMargin=0.7 * inch, leftMargin=0.7 * inch, topMargin=0.7 * inch, bottomMargin=0.7 * inch)
    pdf.build(story)


if __name__ == "__main__":
    main()
