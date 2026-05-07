from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "output" / "01_ColombiaTours_Invoice_Reimbursement_2026-02-03_2027.99.docx"

BLUE = RGBColor(31, 77, 120)
ACCENT = RGBColor(46, 116, 181)
MUTED = RGBColor(96, 106, 116)
LIGHT_FILL = "F2F4F7"
HEADER_FILL = "E8EEF5"


def set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9DEE7"):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360, col_widths=None):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    if col_widths:
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for w in col_widths:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(w))
            grid.append(gc)
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(col_widths[idx]))
                tc_w.set(qn("w:type"), "dxa")


def add_para(doc, text="", size=11, bold=False, italic=False, color=None, align=None, before=0, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def fill_cell(cell, text, bold=False, color=None, size=10.5, align=None):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, 9360, [2520, 6840])
    set_table_borders(table)
    for idx, (label, value) in enumerate(rows):
        set_cell_shading(table.rows[idx].cells[0], LIGHT_FILL)
        fill_cell(table.rows[idx].cells[0], label, bold=True, color=BLUE)
        fill_cell(table.rows[idx].cells[1], value)
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.text = "COLOMBIATOURS.TRAVEL S.A.S. | Invoice / Reimbursement Statement"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    add_para(doc, "COLOMBIATOURS.TRAVEL S.A.S.", size=18, bold=True, color=BLUE, after=1)
    add_para(doc, "Invoice / Reimbursement Statement", size=13, color=MUTED, after=10)

    top = doc.add_table(rows=1, cols=2)
    top.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(top, 9360, [5100, 4260])
    set_table_borders(top)
    left, right = top.rows[0].cells
    set_cell_shading(left, LIGHT_FILL)
    set_cell_shading(right, LIGHT_FILL)
    fill_cell(
        left,
        "Issuer: COLOMBIATOURS.TRAVEL S.A.S.\nNIT: 901346505-1\nAddress: Carrera 17 Nro. 9 - 70 LC 2 Edificio Camellando, Pinares, Pereira, Risaralda, Colombia\nEmail: comunicaciones@colombiatours.travel",
        size=10,
    )
    fill_cell(
        right,
        "Statement No.: CT-WP-2026-0203\nIssue Date: May 6, 2026\nReference Transaction Date: February 03, 2026\nCurrency: USD",
        size=10,
    )

    add_para(doc, "", after=4)
    add_kv_table(
        doc,
        [
            ("Bill To", "Weppa Group LLC"),
            ("Billing Address", "8206 Louisiana Boulevard Northeast, Ste A, Albuquerque, NM 87113, United States"),
            ("Relationship", "Related company under common beneficial ownership by Yeison Gómez"),
            ("Platform", "Bukeer travel technology platform"),
            ("Mercury Transaction Reference", "USD 2,027.99 at COLOMBIATOURS on February 03, 2026 (PT)"),
        ],
    )

    add_para(doc, "Line Items", size=13, bold=True, color=ACCENT, before=8, after=4)
    items = doc.add_table(rows=2, cols=4)
    items.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(items, 9360, [900, 5460, 1500, 1500])
    set_table_borders(items)
    headers = ["Qty", "Description", "Unit Amount", "Line Total"]
    for idx, h in enumerate(headers):
        set_cell_shading(items.rows[0].cells[idx], HEADER_FILL)
        fill_cell(items.rows[0].cells[idx], h, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(items.rows[1].cells[0], "1", align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(
        items.rows[1].cells[1],
        "Intercompany services and reimbursement related to travel industry know-how, supplier coordination, operational requirements, and documented confirmed booking expenses supporting the development and operation of the Bukeer platform.",
        size=10,
    )
    fill_cell(items.rows[1].cells[2], "$2,027.99", align=WD_ALIGN_PARAGRAPH.RIGHT)
    fill_cell(items.rows[1].cells[3], "$2,027.99", align=WD_ALIGN_PARAGRAPH.RIGHT)

    total = doc.add_table(rows=3, cols=2)
    total.alignment = WD_TABLE_ALIGNMENT.RIGHT
    set_table_width(total, 4200, [2100, 2100])
    set_table_borders(total)
    for idx, (label, value, bold) in enumerate(
        [
            ("Subtotal", "$2,027.99", False),
            ("Tax / VAT", "$0.00", False),
            ("Total Due / Reimbursed", "$2,027.99", True),
        ]
    ):
        if bold:
            set_cell_shading(total.rows[idx].cells[0], HEADER_FILL)
            set_cell_shading(total.rows[idx].cells[1], HEADER_FILL)
        fill_cell(total.rows[idx].cells[0], label, bold=bold, color=BLUE if bold else None)
        fill_cell(total.rows[idx].cells[1], value, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT, color=BLUE if bold else None)

    add_para(doc, "Business Purpose And Support", size=13, bold=True, color=ACCENT, before=8, after=4)
    add_para(
        doc,
        "This invoice/reimbursement statement documents intercompany services, travel industry know-how, supplier coordination, operational requirements, and documented confirmed booking-related expenses provided by ColombiaTours.travel S.A.S. to Weppa Group LLC in connection with the Bukeer platform and the February 03, 2026 Mercury transaction.",
    )
    add_para(
        doc,
        "The relationship is supported by the Intercompany Services and Platform Agreement between Weppa Group LLC and ColombiaTours.travel S.A.S. and by company registration documentation for both entities.",
    )

    add_para(doc, "Suggested Supporting Attachments", size=13, bold=True, color=ACCENT, before=8, after=4)
    attachments = [
        "Signed Intercompany Services and Platform Agreement.",
        "ColombiaTours.travel S.A.S. Chamber of Commerce certificate and English summary translation.",
        "Weppa Group LLC formation/EIN documentation.",
        "Mercury transaction screenshot or statement showing the February 03, 2026 payment.",
        "Booking, supplier, or operational support connected to this reimbursement where available.",
    ]
    for item in attachments:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        set_run_font(r, size=10.5)

    add_para(doc, "Certification", size=13, bold=True, color=ACCENT, before=6, after=3)
    add_para(
        doc,
        "I certify that this statement is issued to document and reconcile the referenced intercompany transaction and that the business purpose described above is accurate to the best of my knowledge.",
    )

    sig = doc.add_table(rows=4, cols=1)
    sig.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(sig, 5200, [5200])
    set_table_borders(sig)
    for idx, text in enumerate(
        [
            "Authorized Signature: _______________________________",
            "Name: Yeison Gómez Gómez",
            "Title: Legal Representative, ColombiaTours.travel S.A.S.",
            "Date: _______________________________",
        ]
    ):
        fill_cell(sig.rows[idx].cells[0], text)

    footer = section.footer.paragraphs[0]
    footer.text = "ColombiaTours.travel S.A.S. - Statement CT-WP-2026-0203"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
