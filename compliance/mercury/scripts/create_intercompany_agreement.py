from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "output" / "Weppa_ColombiaTours_Intercompany_Agreement.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(96, 106, 116)
LIGHT_FILL = "F2F4F7"


def set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9DEE7")


def set_table_width(table, width_dxa=9360, col_widths=None):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    if col_widths:
        grid = tbl.tblGrid
        if grid is None:
            grid = OxmlElement("w:tblGrid")
            tbl.append(grid)
        for child in list(grid):
            grid.remove(child)
        for w in col_widths:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(w))
            grid.append(grid_col)
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(col_widths[idx]))
                tc_w.set(qn("w:type"), "dxa")


def add_para(doc, text="", style=None, size=11, bold=False, italic=False, color=None, align=None, before=0, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 12, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_clause(doc, number, title, paragraphs):
    add_heading(doc, f"{number}. {title}", level=2)
    for paragraph in paragraphs:
        add_para(doc, paragraph)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.167
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_metadata_table(doc):
    table = doc.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, 9360, [2160, 7200])
    set_table_borders(table)
    rows = [
        ("Document", "Intercompany Services and Platform Agreement"),
        ("Parties", "Weppa Group LLC and ColombiaTours.travel S.A.S."),
        ("Beneficial Owner", "Yeison Gómez"),
        ("Company Record", "New Mexico Business ID 6558305; organized August 6, 2021"),
        ("Platform", "Bukeer travel technology platform"),
        ("Status", "Ready for completion and signature"),
        ("Effective Date", "May 5, 2026"),
    ]
    for idx, (label, value) in enumerate(rows):
        for cell in table.rows[idx].cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
        set_cell_shading(table.rows[idx].cells[0], LIGHT_FILL)
        p0 = table.rows[idx].cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        set_run_font(r0, bold=True, color=DARK_BLUE)
        p1 = table.rows[idx].cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(value)
        set_run_font(r1)


def add_signature_block(doc):
    add_heading(doc, "15. Signatures", level=2)
    add_para(doc, "IN WITNESS WHEREOF, the Parties execute this Agreement as of the date written below.")
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, 9360, [4680, 4680])
    set_table_borders(table)
    cells = table.rows[0].cells
    cells[0].text = "Weppa Group LLC"
    cells[1].text = "ColombiaTours.travel S.A.S."
    for c in cells:
        set_cell_shading(c, LIGHT_FILL)
    data = [
        ("By: _______________________________", "By: _______________________________"),
        ("Name: Yeison Gómez", "Name: Yeison Gómez Gómez"),
        ("Title: Owner / Authorized Representative", "Title: Legal Representative / Owner"),
        ("Date: ______________________________", "Date: ______________________________"),
    ]
    for r_idx, row in enumerate(data, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=120, bottom=120)
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, bold=(row == table.rows[0]), color=DARK_BLUE if row == table.rows[0] else None)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.text = "Intercompany Services and Platform Agreement"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Weppa Group LLC / ColombiaTours.travel S.A.S."
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)

    add_para(doc, "INTERCOMPANY SERVICES AND PLATFORM AGREEMENT", size=22, bold=True, color=RGBColor(0, 0, 0), after=4)
    add_para(doc, "ACUERDO INTERCOMPAÑÍA DE SERVICIOS Y PLATAFORMA", size=14, color=MUTED, after=16)
    add_metadata_table(doc)

    add_heading(doc, "Parties", level=1)
    add_para(
        doc,
        "This Intercompany Services and Platform Agreement is entered into by and between Weppa Group LLC, a limited liability company organized under the laws of the United States, with its principal business address at 8206 Louisiana Boulevard Northeast, Ste A, Albuquerque, NM 87113, United States, represented by Yeison Gómez, in his capacity as owner/authorized representative, hereinafter referred to as “Weppa”; and ColombiaTours.travel S.A.S., a Colombian company organized under the laws of Colombia, with its principal business address at Carrera 17 Nro. 9 - 70 LC 2 Edificio Camellando, Pinares, Pereira, Risaralda, Colombia, represented by Yeison Gómez Gómez, in his capacity as legal representative/owner, hereinafter referred to as “ColombiaTours.”",
    )
    add_para(doc, "Weppa and ColombiaTours may be referred to individually as a “Party” and collectively as the “Parties.”")

    add_clause(
        doc,
        "1",
        "Background",
        [
            "Weppa is a technology company that develops, operates, and supports digital platforms, software products, cloud infrastructure, automation tools, and related technology services.",
            "Weppa has developed and operates Bukeer, a travel technology platform designed to support travel agencies and tourism operators in managing travel products, suppliers, bookings, customer information, website content, and operational workflows.",
            "ColombiaTours is a Colombian travel agency engaged in the sale, coordination, and operation of domestic and international travel services, including tours, accommodations, transportation, activities, and related travel products.",
            "ColombiaTours has industry experience, operational know-how, supplier relationships, market knowledge, tourism product information, and practical business requirements that support the development, testing, improvement, and commercial use of the Bukeer platform.",
            "Both Weppa and ColombiaTours are under common ownership and control by Yeison Gómez, who is the beneficial owner of both entities.",
            "The Parties wish to formalize their commercial relationship regarding the use of Bukeer, technology services, operational support, supplier coordination, business know-how, reimbursement of expenses, and payments related to documented and confirmed travel bookings.",
        ],
    )

    add_clause(
        doc,
        "2",
        "Purpose",
        [
            "The purpose of this Agreement is to establish the terms under which Weppa provides Bukeer platform access, software development, cloud infrastructure, technical support, automation, and related technology services to ColombiaTours; ColombiaTours provides travel industry know-how, operational requirements, supplier coordination, authorized supplier information, tourism product knowledge, and practical feedback to support the development and operation of Bukeer; and the Parties document intercompany payments, reimbursements, service charges, and supplier-related obligations arising from their commercial relationship.",
        ],
    )

    add_heading(doc, "3. Services Provided By Weppa", level=2)
    add_para(doc, "Weppa may provide ColombiaTours with the following services:")
    add_bullets(
        doc,
        [
            "Access to and use of the Bukeer travel technology platform.",
            "Software development, maintenance, updates, and technical improvements.",
            "Cloud infrastructure, deployment, hosting coordination, and platform monitoring.",
            "Technical support, troubleshooting, automation, and workflow optimization.",
            "Website, booking, content, supplier, and product management tools.",
            "Integration support with payment processors, suppliers, APIs, and operational systems.",
            "Consulting related to digital transformation, sales operations, and technology implementation for travel businesses.",
            "Operational support connected to platform-managed bookings and supplier workflows.",
        ],
    )

    add_heading(doc, "4. Contributions And Services Provided By ColombiaTours", level=2)
    add_para(doc, "ColombiaTours may provide Weppa with the following:")
    add_bullets(
        doc,
        [
            "Travel industry know-how and business expertise.",
            "Operational requirements for the development and improvement of Bukeer.",
            "Feedback based on real travel agency workflows and customer booking processes.",
            "Supplier coordination and authorized supplier-related information.",
            "Tourism product knowledge, destination information, and market requirements.",
            "Testing, validation, and practical use cases for the Bukeer platform.",
            "Support in identifying operational needs related to bookings, suppliers, payments, reservations, and customer service.",
            "Documentation or reports related to confirmed travel operations, where necessary and legally permitted.",
        ],
    )
    add_para(doc, "ColombiaTours shall not provide confidential third-party data to Weppa unless it has the right or authorization to do so.")

    add_clause(
        doc,
        "5",
        "Intercompany Payments And Reimbursements",
        [
            "The Parties acknowledge that, as part of their commercial relationship, payments may be made between Weppa and ColombiaTours for platform access and technology services; software development, support, infrastructure, and consulting; travel industry know-how and operational support; reimbursement of expenses incurred in connection with documented and confirmed travel bookings; supplier-related obligations connected to platform-supported travel operations; and advances or reimbursements for operating costs, provided such amounts are documented by invoice, receipt, statement, or other reasonable business record.",
            "Each intercompany payment shall be supported, when applicable, by invoices, reimbursement requests, receipts, supplier statements, booking confirmations, account statements, or other documentation reasonably sufficient to evidence the business purpose of the payment.",
            "Payments made by Weppa to ColombiaTours may include reimbursements for expenses incurred by ColombiaTours in connection with bookings, supplier coordination, operating obligations, or travel services managed through or supported by Bukeer.",
            "Payments made by ColombiaTours to Weppa may include fees for technology services, software development, platform access, cloud infrastructure, and technical or operational support.",
        ],
    )

    add_clause(
        doc,
        "6",
        "Supplier Payments",
        [
            "The Parties acknowledge that certain travel suppliers, including accommodation providers, wholesalers, reservation platforms, activity operators, and travel service providers, may be paid in connection with confirmed customer bookings.",
            "Where commercially necessary, Weppa may pay suppliers directly or reimburse ColombiaTours for supplier-related obligations connected to bookings managed through the Bukeer/ColombiaTours operational relationship.",
            "Supplier payments shall be supported, where available, by supplier invoices, receipts, booking confirmations, vouchers, account statements, reservation records, or platform records.",
            "The Parties agree that supplier payments are made for legitimate commercial travel operations of the Parties and not for the purpose of providing financial services, escrow, money transmission, or payment processing services to unrelated third parties.",
        ],
    )

    add_clause(
        doc,
        "7",
        "Source And Use Of Funds",
        [
            "The Parties acknowledge that funds used in connection with this Agreement may originate from lawful business revenue, including Stripe or other payment processor payouts, technology service revenue, travel-related sales revenue, platform-related commercial operations, intercompany payments between the Parties, and other lawful business revenue.",
            "Funds may be used for payment of travel suppliers; reimbursement of documented operating expenses; payment of technology, software, cloud, and consulting services; documented and confirmed travel booking obligations; and business operating expenses of the Parties.",
        ],
    )

    add_clause(
        doc,
        "8",
        "Ownership And Intellectual Property",
        [
            "Weppa shall retain ownership of Bukeer, including its software, source code, platform architecture, designs, workflows, documentation, systems, trademarks, and related intellectual property, unless otherwise agreed in writing.",
            "ColombiaTours shall retain ownership of its own brand, customer relationships, travel products, supplier relationships, operational records, and business information.",
            "ColombiaTours grants Weppa the right to use ColombiaTours’ operational feedback, business requirements, travel industry know-how, and authorized supplier/product information solely for the purpose of developing, improving, operating, and supporting Bukeer.",
            "Nothing in this Agreement transfers ownership of either Party’s trademarks, customer lists, confidential information, or intellectual property except as expressly stated.",
        ],
    )

    add_clause(
        doc,
        "9",
        "Confidentiality And Data Protection",
        [
            "Each Party shall keep confidential all non-public business, technical, financial, supplier, customer, and operational information received from the other Party.",
            "Each Party shall use confidential information only for purposes related to this Agreement.",
            "The Parties shall process customer, supplier, and operational data in accordance with applicable data protection laws and only where they have a legitimate business purpose or authorization to do so.",
            "Sensitive customer or supplier information shall not be shared unnecessarily and may be redacted in documents provided to third parties where appropriate.",
        ],
    )

    add_clause(
        doc,
        "10",
        "Records And Documentation",
        [
            "The Parties shall maintain commercially reasonable records supporting intercompany payments, reimbursements, supplier payments, booking-related expenses, technology services, and platform operations.",
            "Such records may include invoices, receipts, contracts, supplier statements, booking confirmations, payment processor reports, bank statements, accounting records, and platform reports.",
            "The Parties may provide such records to banks, payment processors, regulators, auditors, accountants, or legal advisors when reasonably necessary to support compliance, tax, accounting, or business requirements.",
        ],
    )

    add_clause(
        doc,
        "11",
        "Independent Legal Entities",
        [
            "Weppa and ColombiaTours are separate legal entities.",
            "Nothing in this Agreement shall be interpreted as creating a partnership, joint venture, employment relationship, agency relationship, fiduciary relationship, or regulated financial services arrangement between the Parties.",
            "Each Party remains responsible for its own taxes, licenses, regulatory obligations, employees, contractors, and business operations.",
        ],
    )

    add_clause(doc, "12", "Term", ["This Agreement shall become effective as of May 5, 2026 and shall remain in effect until terminated by either Party upon written notice. The Parties acknowledge that this Agreement formalizes and documents an existing commercial and operational relationship between them."])
    add_clause(doc, "13", "Termination", ["Either Party may terminate this Agreement by providing written notice to the other Party. Termination shall not affect obligations incurred before the termination date, including payment, reimbursement, confidentiality, recordkeeping, or supplier-related obligations."])
    add_clause(doc, "14", "Governing Law", ["This Agreement shall be governed by the laws of the State of New Mexico, United States, without regard to conflict of law principles, without prejudice to mandatory Colombian law applicable to ColombiaTours.travel S.A.S. The Parties may mutually agree to resolve disputes through good faith negotiation before initiating legal proceedings."])

    add_signature_block(doc)

    add_heading(doc, "Appendix A - Suggested Invoice Concept", level=1)
    add_para(
        doc,
        "Intercompany services and reimbursement related to travel industry know-how, supplier coordination, operational requirements, and confirmed booking expenses supporting the development and operation of the Bukeer platform.",
        italic=True,
    )

    add_heading(doc, "Appendix B - Documentation Examples", level=1)
    add_bullets(
        doc,
        [
            "Intercompany invoices or reimbursement requests.",
            "Supplier invoices, receipts, vouchers, or reservation confirmations.",
            "Stripe payout summaries and payment processor reports.",
            "Company registration and beneficial ownership documents.",
            "Bukeer platform records connecting bookings, suppliers, and operational payments.",
        ],
    )

    add_heading(doc, "Appendix C - Weppa Group LLC Documentation Summary", level=1)
    add_para(
        doc,
        "Based on available formation and IRS documentation, Weppa Group LLC was organized in New Mexico on August 6, 2021, under Business ID 6558305. The Articles of Organization identify the stated purpose as TECH CONSULTING AND TRAVEL ADVISOR and indicate that the entity is a single-member limited liability company.",
    )
    summary = doc.add_table(rows=5, cols=2)
    summary.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(summary, 9360, [2520, 6840])
    set_table_borders(summary)
    summary_rows = [
        ("Entity Name", "Weppa Group LLC"),
        ("Jurisdiction", "New Mexico, United States"),
        ("Business ID / Formation Date", "6558305 / August 6, 2021"),
        ("Purpose In Articles", "TECH CONSULTING AND TRAVEL ADVISOR"),
        ("Formation/EIN Address On File", "2105 Vista Oeste NW Ste E 1188, Albuquerque, NM 87120, United States"),
    ]
    for idx, (label, value) in enumerate(summary_rows):
        for cell in summary.rows[idx].cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
        set_cell_shading(summary.rows[idx].cells[0], LIGHT_FILL)
        p0 = summary.rows[idx].cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        set_run_font(r0, bold=True, color=DARK_BLUE)
        p1 = summary.rows[idx].cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(value)
        set_run_font(r1)

    add_heading(doc, "Appendix D - ColombiaTours.travel S.A.S. Documentation Summary", level=1)
    add_para(
        doc,
        "Based on the Certificate of Existence and Legal Representation issued by the Chamber of Commerce of Pereira on March 28, 2026, ColombiaTours.travel S.A.S. is a Colombian simplified stock company domiciled in Pereira, Risaralda. The certificate identifies Yeison Gómez Gómez as Gerente and legal representative.",
    )
    ct = doc.add_table(rows=9, cols=2)
    ct.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(ct, 9360, [2520, 6840])
    set_table_borders(ct)
    ct_rows = [
        ("Legal Name", "COLOMBIATOURS.TRAVEL S.A.S."),
        ("Trade Name", "COLOMBIATOURS.TRAVEL"),
        ("NIT", "901346505-1"),
        ("Domicile", "Pereira, Risaralda, Colombia"),
        ("Registration / Incorporation", "Matricula No. 18170618 / December 6, 2019"),
        ("Renewal", "Renewed for 2026 on March 28, 2026"),
        ("Principal Address", "Carrera 17 Nro. 9 - 70 LC 2 Edificio Camellando, Pinares, Pereira, Risaralda"),
        ("Legal Representative", "Yeison Gómez Gómez, C.C. No. 1.088.265.628"),
        ("Main Activities", "Travel agencies, tour operators, reservation services, and related activities"),
    ]
    for idx, (label, value) in enumerate(ct_rows):
        for cell in ct.rows[idx].cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
        set_cell_shading(ct.rows[idx].cells[0], LIGHT_FILL)
        p0 = ct.rows[idx].cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        set_run_font(r0, bold=True, color=DARK_BLUE)
        p1 = ct.rows[idx].cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(value)
        set_run_font(r1)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
