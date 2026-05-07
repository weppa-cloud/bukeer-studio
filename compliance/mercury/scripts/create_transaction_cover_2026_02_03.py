from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "output" / "01_COLOMBIATOURS_2026-02-03_2027.99_Cover_Page.docx"

BLUE = RGBColor(31, 77, 120)
ACCENT = RGBColor(46, 116, 181)
MUTED = RGBColor(96, 106, 116)
LIGHT_FILL = "F2F4F7"
HEADER_FILL = "E8EEF5"


def set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9DEE7"):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360, col_widths=None):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    if col_widths:
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for w in col_widths:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(w))
            grid.append(grid_col)
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(col_widths[idx]))
                tc_w.set(qn("w:type"), "dxa")


def add_para(doc, text="", size=11, bold=False, italic=False, color=None, align=None, before=0, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def fill_cell(cell, text, bold=False, color=None, size=10.5, align=None):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, 9360, [2700, 6660])
    set_table_borders(table)
    for idx, (label, value) in enumerate(rows):
        set_cell_shading(table.rows[idx].cells[0], LIGHT_FILL)
        fill_cell(table.rows[idx].cells[0], label, bold=True, color=BLUE)
        fill_cell(table.rows[idx].cells[1], value)
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.text = "Weppa Group LLC / ColombiaTours.travel S.A.S. | Mercury Transaction Support"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    add_para(doc, "TRANSACTION SUPPORT COVER PAGE", size=20, bold=True, color=BLUE, after=2)
    add_para(doc, "COLOMBIATOURS - USD 2,027.99 - February 03, 2026", size=14, color=MUTED, after=14)

    add_kv_table(
        doc,
        [
            ("Mercury Request", "Relationship and supporting documentation for transaction(s) with $2,027.99 at COLOMBIATOURS on February 03, 2026 (PT)."),
            ("Account Holder", "Weppa Group LLC"),
            ("Counterparty", "ColombiaTours.travel S.A.S."),
            ("Amount / Date", "USD 2,027.99 / February 03, 2026 (PT)"),
            ("Beneficial Owner", "Yeison Gómez is the beneficial owner of Weppa Group LLC and the legal representative/owner of ColombiaTours.travel S.A.S."),
            ("Business Relationship", "ColombiaTours.travel S.A.S. is a related Colombian travel agency under common ownership. ColombiaTours uses and supports Bukeer, Weppa Group LLC's travel technology platform."),
            ("Purpose Of Payment", "Intercompany operations, travel industry know-how, supplier coordination, operational requirements, reimbursements, and documented confirmed booking-related expenses connected to Bukeer/ColombiaTours operations."),
        ],
    )

    add_para(doc, "Summary Explanation", size=13, bold=True, color=ACCENT, before=14, after=5)
    add_para(
        doc,
        "Weppa Group LLC is a U.S. technology company that develops and operates Bukeer, a travel technology platform. ColombiaTours.travel S.A.S. is a Colombian travel agency related to Weppa through common ownership by Yeison Gómez. ColombiaTours provides travel industry know-how, operational requirements, supplier coordination, and practical travel agency workflows that support the development and operation of Bukeer.",
    )
    add_para(
        doc,
        "The referenced payment documents an intercompany reimbursement/service relationship connected to documented travel operations and confirmed bookings. The payment is not related to financial services, escrow, money transmission, or payment processing services for unrelated third parties.",
    )

    add_para(doc, "Documents Included In This Package", size=13, bold=True, color=ACCENT, before=12, after=5)
    docs = [
        "Invoice / reimbursement statement for USD 2,027.99.",
        "Signed Intercompany Services and Platform Agreement between Weppa Group LLC and ColombiaTours.travel S.A.S.",
        "ColombiaTours.travel S.A.S. Chamber of Commerce certificate and English summary translation.",
        "Weppa Group LLC formation/EIN documentation.",
        "Mercury transaction proof for the February 03, 2026 payment.",
        "Booking, supplier, or operational support related to this reimbursement where available.",
    ]
    for item in docs:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=10.5)

    add_para(doc, "Reviewer Note", size=13, bold=True, color=ACCENT, before=12, after=5)
    add_para(
        doc,
        "This package is provided to explain the relationship between Weppa Group LLC, ColombiaTours.travel S.A.S., Bukeer platform operations, and the specific transaction requested by Mercury. Original Spanish corporate documentation is included together with an English convenience summary where applicable.",
        italic=True,
        color=MUTED,
    )

    footer = section.footer.paragraphs[0]
    footer.text = "Transaction Support - COLOMBIATOURS - February 03, 2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
