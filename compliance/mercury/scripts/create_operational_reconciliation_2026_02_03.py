from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "output" / "01_COLOMBIATOURS_Operational_Reconciliation_2026-02-03_2027.99.docx"

NAVY = RGBColor(11, 37, 69)
TEAL = RGBColor(0, 118, 128)
GRAY = RGBColor(93, 104, 116)
LIGHT_TEAL = "E8F4F5"
LIGHT_GRAY = "F5F7FA"


def set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=150, bottom=120, end=150):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D8DEE8"):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360, col_widths=None):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    if col_widths:
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for w in col_widths:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(w))
            grid.append(grid_col)
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(col_widths[idx]))
                tc_w.set(qn("w:type"), "dxa")


def add_para(doc, text="", size=10.5, bold=False, italic=False, color=None, align=None, before=0, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def fill_cell(cell, text, size=10, bold=False, color=None, align=None):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)

    header = section.header.paragraphs[0]
    header.text = "Operational Reconciliation Statement"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color=GRAY)

    add_para(doc, "OPERATIONAL RECONCILIATION STATEMENT", size=18, bold=True, color=NAVY, after=2)
    add_para(doc, "COLOMBIATOURS - USD 2,027.99 - February 03, 2026", size=12.5, color=TEAL, after=14)

    summary = doc.add_table(rows=4, cols=2)
    summary.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(summary, 9360, [2600, 6760])
    set_table_borders(summary)
    rows = [
        ("Prepared For", "Mercury information request regarding transaction at COLOMBIATOURS"),
        ("Prepared By", "ColombiaTours.travel S.A.S. / Weppa Group LLC"),
        ("Transaction Reference", "USD 2,027.99 at COLOMBIATOURS on February 03, 2026 (PT)"),
        ("Purpose", "Reconciliation of intercompany operations, supplier coordination, platform know-how, and documented booking-related expenses connected to Bukeer/ColombiaTours operations."),
    ]
    for idx, (label, value) in enumerate(rows):
        set_cell_shading(summary.rows[idx].cells[0], LIGHT_TEAL)
        fill_cell(summary.rows[idx].cells[0], label, bold=True, color=NAVY)
        fill_cell(summary.rows[idx].cells[1], value)

    add_para(doc, "Statement", size=13, bold=True, color=NAVY, before=14, after=5)
    add_para(
        doc,
        "This statement reconciles the USD 2,027.99 intercompany payment referenced by Mercury on February 03, 2026. The payment relates to ColombiaTours.travel S.A.S. operational support, supplier coordination, travel industry know-how, and documented booking-related obligations connected to the Bukeer travel technology platform operated by Weppa Group LLC.",
    )
    add_para(
        doc,
        "The allocation below is provided as a business reconciliation for the referenced transaction. Supporting records may include invoices, booking records, supplier receipts, Bukeer operational records, payment processor reports, and company documentation.",
    )

    add_para(doc, "Reconciliation Allocation", size=13, bold=True, color=NAVY, before=12, after=5)
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, 9360, [1700, 3960, 2400, 1300])
    set_table_borders(table)
    headers = ["Category", "Description", "Support Type", "Amount"]
    for idx, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], LIGHT_TEAL)
        fill_cell(table.rows[0].cells[idx], h, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)

    data = [
        (
            "Operations",
            "Intercompany services, supplier coordination, operational support, platform know-how, and documented booking-related reconciliation connected to Bukeer/ColombiaTours operations.",
            "Invoice / reimbursement statement; intercompany agreement; company records; booking/supplier support where available",
            "USD 2,027.99",
        ),
        (
            "Total",
            "Total reconciled amount for the February 03, 2026 Mercury transaction.",
            "Invoice / reimbursement statement",
            "USD 2,027.99",
        ),
    ]
    for r_idx, row in enumerate(data, start=1):
        for c_idx, value in enumerate(row):
            if r_idx == len(data):
                set_cell_shading(table.rows[r_idx].cells[c_idx], LIGHT_GRAY)
            align = WD_ALIGN_PARAGRAPH.RIGHT if c_idx == 3 else None
            fill_cell(table.rows[r_idx].cells[c_idx], value, bold=(r_idx == len(data)), color=NAVY if r_idx == len(data) else None, align=align)

    add_para(doc, "Compliance Clarification", size=13, bold=True, color=NAVY, before=12, after=5)
    add_para(
        doc,
        "This reconciliation relates to legitimate commercial operations of related entities under common ownership. It is not intended to evidence financial services, escrow activity, money transmission, or payment processing services for unrelated third parties.",
        italic=True,
        color=GRAY,
    )

    add_para(doc, "Certification", size=13, bold=True, color=NAVY, before=12, after=5)
    add_para(
        doc,
        "I certify that this reconciliation is provided to explain the business purpose of the referenced intercompany transaction and is accurate to the best of my knowledge based on available business records.",
    )
    sig = doc.add_table(rows=4, cols=1)
    sig.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(sig, 5600, [5600])
    set_table_borders(sig)
    for idx, text in enumerate(
        [
            "Authorized Signature: _______________________________",
            "Name: Yeison Gómez Gómez",
            "Title: Legal Representative, ColombiaTours.travel S.A.S.",
            "Date: _______________________________",
        ]
    ):
        fill_cell(sig.rows[idx].cells[0], text)

    footer = section.footer.paragraphs[0]
    footer.text = "Operational Reconciliation - COLOMBIATOURS - February 03, 2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=9, color=GRAY)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
